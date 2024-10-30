from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from .models import Patient, Appointment
from .forms import PatientForm, AppointmentForm, PrescriptionForm
from django.core.paginator import Paginator
from django.http import HttpResponse
from docx import Document
from io import BytesIO
from datetime import datetime



def landing_page(request):
    return render(request, 'landing.html')

@login_required
def doctor_dashboard(request):
    doctor = request.user.doctor  # assuming the logged-in user is a doctor
    patients = Patient.objects.filter(doctor=doctor)

    # Paginate patients list to show 8 per page
    paginator = Paginator(patients, 8)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    appointments = Appointment.objects.all()[:8]  # Display 8 appointments max


    return render(request, 'doctor_dashboard.html', {'page_obj': page_obj, 'appointments': appointments})

@login_required
def add_patient(request):
    if request.method == 'POST':
        form = PatientForm(request.POST)
        if form.is_valid():
            patient = form.save(commit=False)
            patient.doctor = request.user.doctor  # Link to logged-in doctor
            patient.save()
            return redirect('doctor_dashboard')
    else:
        form = PatientForm()
    return render(request, 'add_patient.html', {'form': form})

@login_required
def edit_patient(request, pk):
    patient = get_object_or_404(Patient, pk=pk)
    if request.method == 'POST':
        form = PatientForm(request.POST, instance=patient)
        if form.is_valid():
            form.save()
            return redirect('doctor_dashboard')
    else:
        form = PatientForm(instance=patient)
    return render(request, 'edit_patient.html', {'form': form})

@login_required
def delete_patient(request, pk):
    patient = get_object_or_404(Patient, pk=pk)
    if request.method == 'POST':
        patient.delete()
        return redirect('doctor_dashboard')
    return render(request, 'delete_patient.html', {'patient': patient})

@login_required
def add_appointment(request):
    form = AppointmentForm(request.POST or None)
    if request.method == 'POST' and form.is_valid():
        form.save()
        return redirect('doctor_dashboard')
    return render(request, 'add_appointment.html', {'form': form})

@login_required
def edit_appointment(request, appointment_id):
    appointment = get_object_or_404(Appointment, id=appointment_id)
    form = AppointmentForm(request.POST or None, instance=appointment)
    if request.method == 'POST' and form.is_valid():
        form.save()
        return redirect('doctor_dashboard')
    return render(request, 'edit_appointment.html', {'form': form})

@login_required
def delete_appointment(request, appointment_id):
    appointment = get_object_or_404(Appointment, id=appointment_id)
    if request.method == 'POST':
        appointment.delete()
        return redirect('doctor_dashboard')
    return render(request, 'delete_appointment.html', {'appointment': appointment})

from django.http import HttpResponse
from docx import Document
from io import BytesIO
from datetime import datetime

def create_prescription(request):
    if request.method == 'POST':
        form = PrescriptionForm(request.POST)
        if form.is_valid():
            # Get the form data
            patient_name = form.cleaned_data['patient_name']
            medicines = form.cleaned_data['medicines']
            doctor_notes = form.cleaned_data['doctor_notes']

            # Create a new Word document
            doc = Document()

            # Add some basic info about the medical office
            doc.add_heading('Medical Prescription', 0)
            doc.add_paragraph('Medical Office: Your Medical Office Name')
            doc.add_paragraph('Doctor: Dr. Your Name')
            doc.add_paragraph('Address: Your Address')
            doc.add_paragraph('Phone: Your Phone Number')
            doc.add_paragraph('\n')

            # Add patient info
            doc.add_heading(f'Patient: {patient_name}', level=1)
            doc.add_paragraph(f'Date: {datetime.now().strftime("%Y-%m-%d")}\n')

            # Add medicines
            doc.add_heading('Medicines:', level=2)
            doc.add_paragraph(medicines)

            # Add doctor notes if provided
            if doctor_notes:
                doc.add_heading('Doctor Notes:', level=2)
                doc.add_paragraph(doctor_notes)

            # Prepare the document for download
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            # Send the document as a response
            response = HttpResponse(buffer.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = f'attachment; filename=prescription_{patient_name}.docx'

            return response
    else:
        form = PrescriptionForm()

    return render(request, 'create_prescription.html', {'form': form})
